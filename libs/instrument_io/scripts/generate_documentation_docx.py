
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_documentation():
    # Setup paths
    base_path = Path(r"C:\Users\austi\PROJECTS\UC Irvine\Celia Louise Braun Faiola - FaiolaLab\Notebooks\Emily Truong Notebook\Chemical Inventory and Standards")
    output_path = base_path / "Data_Processing_Documentation.docx"
    
    # Create Document
    doc = Document()
    
    # --- Styles ---
    # Title
    title = doc.add_heading('Chemical Inventory & Standards Consolidation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle/Date
    p = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This document outlines the automated processes used to consolidate, clean, and analyze "
        "the laboratory's chemical inventories and chemical standards. "
        "The goal was to create a unified, high-quality dataset by extracting data from multiple "
        "historical and current Excel files across the lab's OneDrive."
    )

    # --- Section 2: Chemical Inventory Processing ---
    doc.add_heading('2. Chemical Inventory Processing', level=1)
    doc.add_paragraph(
        "We compared the historical 2021 inventory against the current 2025 inventory to identify "
        "discrepancies and 'ghost' items (chemicals present in 2021 but missing in 2025)."
    )
    
    # 2.1 Inputs
    doc.add_heading('2.1 Data Sources', level=2)
    p = doc.add_paragraph(style='List Bullet')
    runner = p.add_run("2021 Source: ")
    runner.bold = True
    p.add_run("Important Docs/Chemical Inventory/02252021-Chemical Inventory.xlsx (Sheets: CiBR-Trac, 428)")
    
    p = doc.add_paragraph(style='List Bullet')
    runner = p.add_run("2025 Source: ")
    runner.bold = True
    p.add_run("Notebooks/Emily Truong Notebook/Chem_Inv.xlsx")

    # 2.2 Methodology
    doc.add_heading('2.2 Processing Methodology', level=2)
    doc.add_paragraph(
        "The data from 2021 was normalized to match the 2025 format. Key steps included:"
    )
    
    steps = [
        "Standardizing column names (e.g., 'Chemical_Name' -> 'Chemical Name').",
        "Merging separate sheets ('CiBR-Trac' and '428') into a single unified list.",
        "Comparing normalized chemical names to identify items missing from the 2025 list."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # --- Section 3: Chemical Standards Extraction ---
    doc.add_heading('3. Chemical Standards Extraction', level=1)
    doc.add_paragraph(
        "A comprehensive scan of the lab's digital records was performed to build a master list of "
        "chemical standards. This involved scraping data from calibration logs, project files, and "
        "historical lists."
    )

    # 3.1 Sources
    doc.add_heading('3.1 Scanned Files', level=2)
    doc.add_paragraph("Data was successfully extracted from the following files:")
    
    sources = [
        "Response factors.xlsx",
        "Soil VOC quantitation.xlsx",
        "Standard Calculations (1).xlsx (Avisa)",
        "8mix_calc.xlsx",
        "std_tidy.xlsx",
        "StandardsAndCals.xlsx (parsed from mixture descriptions)",
        "ChiralStandards_Cal - Updated.xlsx",
        "Universal Chemical List.xlsx",
        "Claire Chemical Standard List-Faiola.xlsx",
        "OLD_CompiledStandardList.xlsx"
    ]
    for source in sources:
        doc.add_paragraph(source, style='List Bullet')

    # 3.2 Cleaning Rules
    doc.add_heading('3.2 Data Cleaning & Quality Control', level=2)
    doc.add_paragraph(
        "To ensure the master list contains only valid chemicals, strict filtering rules were applied:"
    )
    
    rules = [
        "Garbage Removal: Generic placeholders (e.g., 'MT1', 'UNKNOWN', 'SQT20', 'Sample') were strictly excluded.",
        "Normalization: Chemical names were standardized (e.g., 'a-pinene' → 'α-Pinene').",
        "Deduplication: Exact duplicates across multiple files were removed, keeping the most relevant entry.",
        "Mixture Parsing: Text descriptions of mixtures (e.g., 'a-pinene / linalool') were split to list individual components."
    ]
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')

    # --- Section 4: Output Files ---
    doc.add_heading('4. Generated Reports', level=1)
    doc.add_paragraph(
        "The following Excel files have been generated and saved in the 'Chemical Inventory and Standards' folder:"
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Name'
    hdr_cells[1].text = 'Description'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    outputs = [
        ("Chemical_Inventory_List_2021.xlsx", "Cleaned, unified copy of the 2021 inventory."),
        ("Chemical_Inventory_List_2025.xlsx", "Cleaned copy of the current 2025 inventory."),
        ("Chemical_Inventory_Gap_Analysis.xlsx", "List of items present in 2021 but missing in 2025."),
        ("Chemical_Standards_List_2025.xlsx", "Master list of 88+ unique standards extracted from all scanned files.")
    ]

    for filename, desc in outputs:
        row_cells = table.add_row().cells
        row_cells[0].text = filename
        row_cells[1].text = desc

    # Save
    doc.save(output_path)
    print(f"Documentation created at: {output_path}")

if __name__ == "__main__":
    create_documentation()
